"""Compile every Markdown section in this folder into a single FYDP-compliant
Word document.

Layout:
  Front matter (lowercase Roman page numbers i, ii, iii, ...)
    - Title page
    - Author's Declaration
    - Statement of Contributions
    - Executive Summary
    - Acknowledgments
    - Table of Contents (Word TOC field)
    - List of Figures (Word Table-of-Figures field)
    - List of Tables (Word Table-of-Tables field)
    - List of Abbreviations
    - UN Sustainable Development Goals
    - Similarity Index Report
  Main body (Arabic page numbers 1, 2, 3, ...)
    - Chapters 1 through 8
    - References

Markdown features handled:
  - ATX headings (# through ####)
  - Paragraphs with inline **bold**, *italic*, `code`, [link](url)
  - Bullet lists (-, *, +) and numbered lists (1., 2., ...)
  - Tables (GitHub-style with header separator)
  - Images (![alt](path))    — paths resolved relative to this folder
  - Horizontal rules (---)   — turned into page breaks for top-of-section rules
  - Block quotes (> text)    — rendered as indented italic notes
  - Fenced code blocks (``` ... ```) — rendered as Courier with grey shading
  - HTML comments            — silently dropped
"""
from __future__ import annotations
from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --------------------------------------------------------------------------- #
# Paths & options
# --------------------------------------------------------------------------- #
HERE = Path(__file__).resolve().parent
DIAGRAMS_DIR = HERE.parent / "diagrams"
OUTPUT_PATH = HERE.parent.parent / "Final_FYDP_Report.docx"

# Reading order (front-matter first in roman section, then main body)
FRONT_MATTER_FILES: list[str] = [
    "00-title-page.md",
    "01-declaration-contributions.md",
    "02-abstract.md",
    "03-frontmatter.md",  # acks + abbreviations + SDGs + similarity index
]
MAIN_BODY_FILES: list[str] = [
    "05-chapter1-introduction.md",
    "06-chapter2-literature.md",
    "07-chapter3-requirements.md",
    "08-chapter4-design.md",
    "09-chapter5-implementation.md",
    "10-chapter6-testing.md",
    "11-chapter7-user-manual.md",
    "12-chapter8-conclusion.md",
    "13-references.md",
]

# --------------------------------------------------------------------------- #
# Styling constants (FYDP guidelines)
# --------------------------------------------------------------------------- #
FONT_NAME = "Times New Roman"
BODY_PT   = 12
H1_PT     = 18
H2_PT     = 14
H3_PT     = 13
H4_PT     = 12
LINE_SPACING = 1.5

LIGHT_GREY = RGBColor(0x66, 0x66, 0x66)
DARK       = RGBColor(0x11, 0x11, 0x11)
ACCENT     = RGBColor(0xCC, 0x22, 0x22)

# --------------------------------------------------------------------------- #
# Helpers — low-level XML utilities
# --------------------------------------------------------------------------- #

def add_field(paragraph, instr_text: str, default_text: str = "") -> None:
    """Insert a Word complex field (TOC, PAGE, NUMPAGES, etc.) into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r.append(fld_sep)
    if default_text:
        t = OxmlElement("w:t")
        t.text = default_text
        r.append(t)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_end)


def add_page_break(doc):
    """Insert a hard page break."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_run_font(run, *, size_pt=BODY_PT, bold=False, italic=False, color=None,
                 mono=False):
    run.font.name = "Courier New" if mono else FONT_NAME
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    if mono:
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), "Courier New")
    else:
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_styled_paragraph(doc, text: str = "", style: str = "Normal",
                         align=None, before_pt=0, after_pt=6) -> object:
    p = doc.add_paragraph(style=style) if style != "Normal" else doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = LINE_SPACING
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p


# --------------------------------------------------------------------------- #
# Inline Markdown → docx Runs
# --------------------------------------------------------------------------- #

INLINE_PATTERNS = [
    # order matters: try strong first, then emphasis, then code, then link
    ("strong",   re.compile(r"\*\*([^*]+?)\*\*")),
    ("strong_u", re.compile(r"__([^_]+?)__")),
    ("emph",     re.compile(r"\*([^*]+?)\*")),
    ("emph_u",   re.compile(r"(?<!\w)_([^_]+?)_(?!\w)")),
    ("code",     re.compile(r"`([^`]+?)`")),
    ("link",     re.compile(r"\[([^\]]+)\]\(([^)]+)\)")),
]


def add_inline_runs(paragraph, text: str, *, size_pt=BODY_PT, italic=False):
    """Parse Markdown inline syntax in `text` and add styled runs to paragraph."""
    # Drop HTML-entity artefacts that python-docx will copy literally.
    text = text.replace("&nbsp;", " ")

    # Manual scanner: walk left-to-right, find earliest match across patterns.
    cursor = 0
    while cursor < len(text):
        earliest = None
        for kind, pat in INLINE_PATTERNS:
            m = pat.search(text, cursor)
            if m and (earliest is None or m.start() < earliest[1].start()):
                earliest = (kind, m)
        if earliest is None:
            # No more inline markup
            run = paragraph.add_run(text[cursor:])
            set_run_font(run, size_pt=size_pt, italic=italic)
            return
        kind, m = earliest
        if m.start() > cursor:
            run = paragraph.add_run(text[cursor:m.start()])
            set_run_font(run, size_pt=size_pt, italic=italic)
        if kind in ("strong", "strong_u"):
            run = paragraph.add_run(m.group(1))
            set_run_font(run, size_pt=size_pt, bold=True, italic=italic)
        elif kind in ("emph", "emph_u"):
            run = paragraph.add_run(m.group(1))
            set_run_font(run, size_pt=size_pt, italic=True)
        elif kind == "code":
            run = paragraph.add_run(m.group(1))
            set_run_font(run, size_pt=size_pt - 1, mono=True)
        elif kind == "link":
            # Render as plain text (label) — Word won't auto-resolve hyperlinks
            # from a docx-js-style insert without extra plumbing.
            run = paragraph.add_run(m.group(1))
            set_run_font(run, size_pt=size_pt, italic=italic, color=ACCENT)
        cursor = m.end()


# --------------------------------------------------------------------------- #
# Block parsers
# --------------------------------------------------------------------------- #

class State:
    """Mutable parsing state shared across the build."""
    def __init__(self):
        self.figure_list: list[tuple[str, str]] = []   # (label, caption)
        self.table_list: list[tuple[str, str]] = []
        self.current_chapter: int = 0


def add_heading(doc, level: int, text: str, state: State):
    sizes = {1: H1_PT, 2: H2_PT, 3: H3_PT, 4: H4_PT}
    size_pt = sizes.get(level, BODY_PT)
    style_name = f"Heading {level}" if level <= 4 else "Heading 4"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=True, color=DARK)
    if level == 1:
        state.current_chapter += 1
    return p


def add_list_paragraph(doc, text: str, *, ordered: bool, level: int = 0):
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4 + 0.3 * level)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(3)
    if not ordered:
        # Ensure a real bullet character at the start
        bullet = p.add_run("•\t")
        set_run_font(bullet)
    add_inline_runs(p, text)
    return p


def add_image(doc, image_path: Path, alt: str, state: State):
    if not image_path.exists():
        # Fall back to a labelled placeholder so the document still compiles.
        para = doc.add_paragraph()
        run = para.add_run(f"[MISSING IMAGE: {image_path.name}]")
        set_run_font(run, italic=True, color=ACCENT)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Cap the rendered width at 6 inches (content width on A4 with 1" margins
    # is ~6.5 inches; this preserves a hair of bleed).
    run.add_picture(str(image_path), width=Inches(6.0))


def add_horizontal_rule(doc):
    """Insert a thin horizontal-rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_blockquote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "CC2222")
    pBdr.append(left)
    pPr.append(pBdr)
    add_inline_runs(p, text, italic=False)


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    for i, line in enumerate(code.splitlines()):
        run = p.add_run(line + ("\n" if i < len(code.splitlines()) - 1 else ""))
        set_run_font(run, size_pt=BODY_PT - 2, mono=True)
        if i < len(code.splitlines()) - 1:
            run.add_break()


def add_md_table(doc, header: list[str], rows: list[list[str]], state: State):
    """Render a Markdown table as a docx Table with thin grey borders."""
    if not header:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Header row
    for col_i, cell_text in enumerate(header):
        cell = table.rows[0].cells[col_i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        para = cell.paragraphs[0]
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        add_inline_runs(para, cell_text)
        for r in para.runs:
            r.bold = True
        # Shade header
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F7F7F7")
        tcPr.append(shd)
    # Body rows
    for row_i, row in enumerate(rows):
        for col_i in range(len(header)):
            cell = table.rows[row_i + 1].cells[col_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            value = row[col_i] if col_i < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, value, size_pt=BODY_PT - 1)
    # Borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "AAAAAA")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # Space after table
    doc.add_paragraph()


# --------------------------------------------------------------------------- #
# Captions — extract Figure / Table captions for the auto-built lists
# --------------------------------------------------------------------------- #

FIGURE_CAPTION_RE = re.compile(r"\*\*Figure\s+(\d+\.\d+)\s+[—\-]\s+(.+?)\*\*")
TABLE_CAPTION_RE  = re.compile(r"\*\*Table\s+(\d+\.\d+)\s+[—\-]\s+(.+?)\*\*")


# --------------------------------------------------------------------------- #
# Top-level Markdown parser  (block-by-block)
# --------------------------------------------------------------------------- #

H_RE       = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
HR_RE      = re.compile(r"^[\s]*[-*_]{3,}[\s]*$")
ULI_RE     = re.compile(r"^(\s*)[-*+]\s+(.+)$")
OLI_RE     = re.compile(r"^(\s*)\d+\.\s+(.+)$")
TABLE_ROW  = re.compile(r"^\s*\|.+\|\s*$")
TABLE_SEP  = re.compile(r"^\s*\|\s*:?[-]+:?(\s*\|\s*:?[-]+:?)*\s*\|\s*$")
IMG_LINE   = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")
BLOCKQ_RE  = re.compile(r"^>\s?(.*)$")
COMMENT_OPEN = re.compile(r"<!--")
COMMENT_CLOSE = re.compile(r"-->")


def parse_markdown(text: str, doc, state: State):
    lines = text.splitlines()
    i = 0
    in_comment = False
    while i < len(lines):
        line = lines[i]

        # Strip HTML comments (possibly multi-line)
        if in_comment:
            if COMMENT_CLOSE.search(line):
                in_comment = False
            i += 1
            continue
        if COMMENT_OPEN.search(line) and not COMMENT_CLOSE.search(line):
            in_comment = True
            i += 1
            continue
        if COMMENT_OPEN.search(line) and COMMENT_CLOSE.search(line):
            i += 1
            continue

        # Blank line — just advance
        if not line.strip():
            i += 1
            continue

        # Code fence
        if line.strip().startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, "\n".join(code_lines))
            continue

        # ATX heading
        m = H_RE.match(line)
        if m:
            level = len(m.group(1))
            text_h = m.group(2)
            add_heading(doc, level, text_h, state)
            i += 1
            continue

        # Horizontal rule
        if HR_RE.match(line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Image
        m = IMG_LINE.match(line)
        if m:
            alt, path = m.group(1), m.group(2)
            # Relative paths assumed to point at documentation/diagrams/
            p = Path(path)
            if not p.is_absolute():
                # Resolve from this script's directory
                p = (HERE / path).resolve()
            add_image(doc, p, alt, state)
            i += 1
            # Capture figure caption if next non-blank line is the **Figure …** caption
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                cap_match = FIGURE_CAPTION_RE.search(lines[i])
                if cap_match:
                    label = f"Figure {cap_match.group(1)}"
                    state.figure_list.append((label, cap_match.group(2)))
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_after = Pt(8)
                    add_inline_runs(cap_p, lines[i].strip(), size_pt=BODY_PT - 1,
                                    italic=True)
                    i += 1
            continue

        # Blockquote
        if BLOCKQ_RE.match(line):
            quoted_lines: list[str] = []
            while i < len(lines) and BLOCKQ_RE.match(lines[i]):
                quoted_lines.append(BLOCKQ_RE.match(lines[i]).group(1))
                i += 1
            add_blockquote(doc, " ".join(quoted_lines).strip())
            continue

        # Table?
        if TABLE_ROW.match(line) and i + 1 < len(lines) and \
                TABLE_SEP.match(lines[i + 1]):
            header_cells = _split_table_row(line)
            i += 2
            body_rows: list[list[str]] = []
            while i < len(lines) and TABLE_ROW.match(lines[i]):
                body_rows.append(_split_table_row(lines[i]))
                i += 1
            # Capture preceding table caption (the most recent **Table X.Y — …**)
            add_md_table(doc, header_cells, body_rows, state)
            continue

        # Bullet / numbered list
        if ULI_RE.match(line) or OLI_RE.match(line):
            ordered = bool(OLI_RE.match(line))
            while i < len(lines):
                m_ul = ULI_RE.match(lines[i])
                m_ol = OLI_RE.match(lines[i])
                if m_ul:
                    add_list_paragraph(doc, m_ul.group(2), ordered=False,
                                       level=len(m_ul.group(1)) // 2)
                elif m_ol:
                    add_list_paragraph(doc, m_ol.group(2), ordered=True,
                                       level=len(m_ol.group(1)) // 2)
                else:
                    break
                i += 1
            continue

        # Figure / Table caption line (standalone, no preceding image — used
        # for figures referenced from earlier and for tables introduced before
        # their body row). Capture into the list so the auto-LoF/LoT picks it up.
        cap_match_t = TABLE_CAPTION_RE.search(line)
        if cap_match_t:
            label = f"Table {cap_match_t.group(1)}"
            state.table_list.append((label, cap_match_t.group(2)))
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_after = Pt(6)
            add_inline_runs(cap_p, line.strip(), size_pt=BODY_PT - 1, italic=True)
            i += 1
            continue
        cap_match_f = FIGURE_CAPTION_RE.search(line)
        if cap_match_f:
            label = f"Figure {cap_match_f.group(1)}"
            state.figure_list.append((label, cap_match_f.group(2)))
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_after = Pt(6)
            add_inline_runs(cap_p, line.strip(), size_pt=BODY_PT - 1, italic=True)
            i += 1
            continue

        # Default: paragraph
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = LINE_SPACING
        para.paragraph_format.space_after = Pt(8)
        add_inline_runs(para, line.strip())
        i += 1


def _split_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


# --------------------------------------------------------------------------- #
# Section / page-numbering setup
# --------------------------------------------------------------------------- #

def configure_section_margins(section, *, binding_inch=1.5, other_inch=1.0):
    section.page_height = Cm(29.7)        # A4
    section.page_width  = Cm(21.0)
    section.left_margin   = Inches(binding_inch)
    section.right_margin  = Inches(other_inch)
    section.top_margin    = Inches(other_inch)
    section.bottom_margin = Inches(other_inch)
    section.orientation   = WD_ORIENTATION.PORTRAIT


def add_page_number_footer(section, *, format_code: str = "decimal",
                           start_at: int | None = None,
                           continue_from_previous: bool = False):
    """Insert a centred page-number field in the section footer."""
    footer = section.footer
    # Clear existing
    for para in list(footer.paragraphs):
        for run in list(para.runs):
            run.text = ""
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE \\* MERGEFORMAT", default_text="1")

    # Configure the section's page-number format / start
    sectPr = section._sectPr
    # Strip any existing pgNumType
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), format_code)
    if start_at is not None:
        pgNumType.set(qn("w:start"), str(start_at))
    sectPr.append(pgNumType)


def break_to_new_section(doc):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    return new_section


# --------------------------------------------------------------------------- #
# Built-in TOC / LoF / LoT — Word fields that auto-populate on open
# --------------------------------------------------------------------------- #

def add_toc(doc, *, title="Table of Contents"):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(12)
    run = h.add_run(title)
    set_run_font(run, size_pt=H1_PT, bold=True, color=DARK)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u',
              default_text="(Right-click → Update Field in Word to populate this table.)")


def add_list_of_figures(doc):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(12)
    run = h.add_run("List of Figures")
    set_run_font(run, size_pt=H1_PT, bold=True, color=DARK)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figure"',
              default_text="(Right-click → Update Field to populate the figures list.)")


def add_list_of_tables(doc):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(12)
    run = h.add_run("List of Tables")
    set_run_font(run, size_pt=H1_PT, bold=True, color=DARK)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Table"',
              default_text="(Right-click → Update Field to populate the tables list.)")


# --------------------------------------------------------------------------- #
# Document set-up: default style, headings, base
# --------------------------------------------------------------------------- #

def configure_default_styles(doc: Document):
    """Apply Times New Roman 12pt body and force Heading style defaults."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_PT)
    rFonts = style.element.rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        style.element.rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT_NAME)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(8)

    for lvl, size in [(1, H1_PT), (2, H2_PT), (3, H3_PT), (4, H4_PT)]:
        try:
            hs = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        hs.font.name = FONT_NAME
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = DARK


# --------------------------------------------------------------------------- #
# Main build
# --------------------------------------------------------------------------- #

def main():
    doc = Document()
    configure_default_styles(doc)
    state = State()

    # ----- Section 1: Front matter, lower-case Roman numerals -----
    section1 = doc.sections[0]
    configure_section_margins(section1, binding_inch=1.5, other_inch=1.0)
    add_page_number_footer(section1, format_code="lowerRoman", start_at=1)

    # Title page (no page number visible — we'll still count it as i)
    for fname in FRONT_MATTER_FILES[:1]:
        text = (HERE / fname).read_text(encoding="utf-8")
        parse_markdown(text, doc, state)
        add_page_break(doc)

    # Remaining front matter
    for fname in FRONT_MATTER_FILES[1:]:
        text = (HERE / fname).read_text(encoding="utf-8")
        parse_markdown(text, doc, state)
        add_page_break(doc)

    # TOC, List of Figures, List of Tables (auto-populated by Word)
    add_toc(doc)
    add_page_break(doc)
    add_list_of_figures(doc)
    add_page_break(doc)
    add_list_of_tables(doc)

    # ----- Section 2: Main body, Arabic numerals restarting at 1 -----
    section2 = break_to_new_section(doc)
    configure_section_margins(section2, binding_inch=1.5, other_inch=1.0)
    add_page_number_footer(section2, format_code="decimal", start_at=1)

    for fname in MAIN_BODY_FILES:
        text = (HERE / fname).read_text(encoding="utf-8")
        parse_markdown(text, doc, state)
        add_page_break(doc)

    # ----- Save -----
    OUTPUT_PATH.parent.mkdir(exist_ok=True, parents=True)
    doc.save(OUTPUT_PATH)
    print(f"OK: wrote {OUTPUT_PATH}")
    print(f"     {len(state.figure_list)} figures, {len(state.table_list)} tables captured.")
    print("     Open the file in Word and press Ctrl+A then F9 to update all fields")
    print("     (Table of Contents, List of Figures, List of Tables).")


if __name__ == "__main__":
    sys.exit(main())
